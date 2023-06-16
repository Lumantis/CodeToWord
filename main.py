import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm
from ttkbootstrap import Style, Button, Label
from PIL import Image, ImageTk

# Liste des extensions de fichiers qui vont être lues
EXTENSIONS = ['.js', '.env', '.py', '.css', '.html', '.txt', '.json', '.yaml', '.yml', '.ini']
ERROR_MESSAGE_SELECTION = "Aucun répertoire sélectionné. Veuillez d'abord sélectionner un répertoire."

def lire_fichier(chemin: str) -> str:
    try:
        with open(chemin, "r", encoding='utf-8') as fichier:
            return fichier.read()
    except Exception as e:
        raise ValueError(f"Erreur lors de la lecture du fichier {chemin}: {e}")

def selectionner_repertoire():
    repertoire = filedialog.askdirectory()
    if repertoire:
        app.repertoire = repertoire
        label_repertoire["text"] = f"Répertoire sélectionné : {repertoire}"
        bouton_creer_document["state"] = "normal"

def creer_document():
    if not app.repertoire:
        messagebox.showinfo("Aucun répertoire sélectionné", ERROR_MESSAGE_SELECTION)
        return

    if not os.listdir(app.repertoire):
        messagebox.showinfo("Répertoire vide", "Le répertoire sélectionné est vide. Veuillez sélectionner un répertoire contenant des fichiers.")
        return

    doc = Document()

    # Ajout d'un titre au document
    titre = doc.add_paragraph()
    run = titre.add_run("Fichiers de code :")
    run.font.size = Pt(24)
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    fichiers_a_traiter = [os.path.join(chemin_repertoire, nom_fichier)
                          for chemin_repertoire, _, noms_fichier in os.walk(app.repertoire)
                          for nom_fichier in noms_fichier
                          if any(nom_fichier.endswith(ext) for ext in EXTENSIONS)]

    for chemin_fichier in tqdm(fichiers_a_traiter, desc="Traitement des fichiers", unit="fichier"):
        try:
            contenu_fichier = lire_fichier(chemin_fichier)
        except ValueError as e:
            messagebox.showerror("Erreur", str(e))
            continue

        # Ajout du nom et de l'emplacement du fichier au document
        paragraphe = doc.add_paragraph()
        run = paragraphe.add_run(f"Fichier : {os.path.basename(chemin_fichier)}\nEmplacement : {chemin_fichier}\n")
        run.bold = True
        run.italic = True

        # Ajout du contenu du fichier au document
        doc.add_paragraph(contenu_fichier)

    chemin_fichier = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Document Word", "*.docx")])
    if not chemin_fichier:  # L'utilisateur a annulé la boîte de dialogue de sauvegarde
        return

    try:
        doc.save(chemin_fichier)
        messagebox.showinfo("Succès", "Document créé avec succès !")
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur lors de la sauvegarde du document : {e}")
    finally:
        app.quit()

app = Style().master
app.geometry("600x400")  # Taille de la fenêtre
app.title("Code To Docx")  # Titre de la fenêtre
app.repertoire = None

# Chargement et placement du logo
logo = Image.open("Codetoword.png")
logo = logo.resize((30, 30), Image.LANCZOS)
photo_logo = ImageTk.PhotoImage(logo)
logo_label = Label(image=photo_logo)
logo_label.image = photo_logo
logo_label.pack(pady=10)

label_repertoire = Label(app, text="Aucun répertoire sélectionné")
label_repertoire.pack(padx=10, pady=10)

bouton_selectionner_repertoire = Button(app, text="Sélectionner un répertoire", command=selectionner_repertoire)
bouton_selectionner_repertoire.pack(padx=10, pady=10)

bouton_creer_document = Button(app, text="Créer Document", command=creer_document, state="disabled")
bouton_creer_document.pack(padx=10, pady=10)

# Ajouter un menu
menu = tk.Menu(app)
app.config(menu=menu)

menu_fichier = tk.Menu(menu)
menu.add_cascade(label="Fichier", menu=menu_fichier)
menu_fichier.add_command(label="Quitter", command=app.quit)

app.mainloop()
